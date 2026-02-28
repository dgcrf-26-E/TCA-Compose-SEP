import io
import os
from django.shortcuts import redirect, render
from django.urls import reverse_lazy
from django.conf import settings
from django.http import FileResponse
from django.contrib import messages

from django.views.generic import FormView, TemplateView
from django.contrib.auth.mixins import LoginRequiredMixin

from .mixins import AccessKeyRequiredMixin
from panel.forms import AccessKeyForm

from .models import (
    Formulario,
    Seccion,
    SubSeccion,
    Pregunta,
    RespuestaFormulario,
    RegistroTemporal,
    AccionesTemporal,
    ReporteGenerado,
)
from usuarios.models import Area, Rubro, Periodo, Registro, Acciones
from .serializers import (
    FormularioSerializer,
    SeccionSerializer,
    SubseccionSerializer,
    PreguntaSerializer,
    RespuestaFormularioSerializer,
    RegistroTemporalSerializer,
    AccionesTemporalSerializer,
    AreasSerializer,
    RubrosSerializer,
    PeriodosSerializer,
    ReporteGeneradoSerializer,
)

from django.db.models import Count

from rest_framework import viewsets
from rest_framework.decorators import action
from rest_framework.response import Response

# Create your views here.

class AccessKeyView(FormView):
    template_name = 'panel/access_key.html'
    form_class = AccessKeyForm
    success_url = reverse_lazy('superv_urls:index')

    def form_valid(self, form):
        llave = form.cleaned_data['key']
        if llave == settings.SUPERVISION_KEY:
            # self.request.session['has_access'] = True
            self.request.session['access_keyS'] = llave
            return super().form_valid(form)
        form.add_error('key', 'Clave incorrecta')
        return self.form_invalid(form)

class InicioView(LoginRequiredMixin, AccessKeyRequiredMixin, TemplateView):
    template_name = "supervision/index_s.html"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        req = super().get_context_data(request=self.request.get_host())
        context["debug"] = settings.DEBUG
        context["url_i"] = req['request']
        # print(context["url_i"])
        return context
    
    def get(self, request, *args, **kwargs):
        return super().get(request, *args, **kwargs)


class FormularioViewSet(viewsets.ModelViewSet):
    queryset = Formulario.objects.all()
    serializer_class = FormularioSerializer
    
    @action(detail=False, methods=["get"])
    def ultimo(self, request):
        ultimo = self.get_queryset().order_by("-id").first()
        if not ultimo:
            return Response({}, status=200)
        serializer = self.get_serializer(ultimo)
        return Response(serializer.data)


class SeccionViewSet(viewsets.ModelViewSet):
    queryset = Seccion.objects.all()
    serializer_class = SeccionSerializer


class SubSeccionViewSet(viewsets.ModelViewSet):
    queryset = SubSeccion.objects.all()
    serializer_class = SubseccionSerializer


class PreguntaViewSet(viewsets.ModelViewSet):
    queryset = Pregunta.objects.all()
    serializer_class = PreguntaSerializer


class RespuestaFormularioViewSet(viewsets.ModelViewSet):
    queryset = RespuestaFormulario.objects.all()
    serializer_class = RespuestaFormularioSerializer


class RegistroTemporalViewSet(viewsets.ModelViewSet):
    queryset = RegistroTemporal.objects.all()
    serializer_class = RegistroTemporalSerializer

    def perform_destroy(self, instance):
        # Delete related actions first
        # related_name='accionRTemporal' from AccionesTemporal definitions
        if instance.accionRTemporal.exists():
            for accion in instance.accionRTemporal.all():
                accion.delete()
        instance.delete()


class AccionesTemporalViewSet(viewsets.ModelViewSet):
    queryset = AccionesTemporal.objects.all()
    serializer_class = AccionesTemporalSerializer

class AreasViewSet(viewsets.ModelViewSet):
    queryset = Area.objects.all().order_by("idArea")
    serializer_class = AreasSerializer

class RubrosViewSet(viewsets.ModelViewSet):
    queryset = Rubro.objects.all().order_by("idRubro")
    serializer_class = RubrosSerializer

class PeriodosViewSet(viewsets.ModelViewSet):
    queryset = Periodo.objects.all()
    serializer_class = PeriodosSerializer

class ReporteGeneradoViewSet(viewsets.ModelViewSet):
    queryset = ReporteGenerado.objects.all()
    serializer_class = ReporteGeneradoSerializer
    

def _construir_documento_word():
    from docxtpl import DocxTemplate
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from datetime import datetime

    # Obtener los registros con sus relaciones precargadas
    registros = (
        RegistroTemporal.objects
        .prefetch_related('accionRTemporal', 'area', 'rubro')
        .all()
    )

    # Conteo de rubros único
    conteo_rubros = (
        RegistroTemporal.objects
        .values('rubro__tipo')
        .annotate(total=Count('idRegistro', distinct=True))
        .filter(total__gt=0)
    )

    if registros.exists():
        datosA = []
        for reg in registros:
            acciones = reg.accionRTemporal.all()
            antecedentes = "\n".join([a.antecedente for a in acciones if a.antecedente])
            descripciones = "\n".join([a.descripcion for a in acciones if a.descripcion])
            areas = ", ".join([area.nickname for area in reg.area.all()])
            rubro_tipo = reg.rubro.first().tipo if reg.rubro.exists() else "N/A"

            datosA.append({
                "rubro__tipo": rubro_tipo,
                "claveAcuerdo": reg.claveAcuerdo,
                "accionRTemporal__antecedente": antecedentes,
                "accionRTemporal__descripcion": descripciones,
                "fecha_termino": reg.fecha_termino.strftime("%d/%m/%Y"),
                "area__accionA2Temporal__area2__nickname": areas,
            })

        inicio_f = registros.first().fecha_inicio
        partes_clave = registros.first().claveAcuerdo.split("/")
        oficina_d = partes_clave[1] if len(partes_clave) > 1 else "COLIMA"
        
        context = {
            "fecha": inicio_f.strftime("%d de %B de %Y") if inicio_f else datetime.now().strftime("%d de %B de %Y"),
            "oficina": oficina_d
        }
    else:
        datosA = [
            {"rubro__tipo": "INFO 1", "claveAcuerdo": "001/DEMO/01/2025", "accionRTemporal__antecedente": "Sin antecedentes", "accionRTemporal__descripcion": "Sin descripción", "fecha_termino": "01/01/2025", "area__accionA2Temporal__area2__nickname": "Área Demo",},
        ]
        context = {
            "fecha": datetime.now().strftime("%d de %B de %Y"),
            "oficina": "DEMO"
        }

    ruta_word = os.path.join(settings.BASE_DIR, 'zdata', 'plantilla.docx')
    doc = DocxTemplate(ruta_word)
    doc.render(context)
    
    temp_stream = io.BytesIO()
    doc.save(temp_stream)
    temp_stream.seek(0)

    document = Document(temp_stream)
    table = document.tables[0]

    coloresT1 = {0: "EFE8DE", 1: "FFFFFF"}
    coloresT2 = {0: "AEE0D5", 1: "FFFFFF"}
    alineacion = { -1: WD_ALIGN_PARAGRAPH.LEFT, 0: WD_ALIGN_PARAGRAPH.CENTER, 1: WD_ALIGN_PARAGRAPH.RIGHT, 2: WD_ALIGN_PARAGRAPH.JUSTIFY }

    def modificar_celda(celda, texto="", text_size=10, bg_color_hex="FFFFFF", text_color_hex="000000", bold = False, italic=False, aliniar = 0):
        aliniarV = alineacion[aliniar] if aliniar in alineacion else alineacion[0]
        celda.text = str(texto)
        textColor = RGBColor(int(text_color_hex[0:2],16), int(text_color_hex[2:4],16), int(text_color_hex[4:6],16))
        celda.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in celda.paragraphs:
            paragraph.alignment = aliniarV
            for run in paragraph.runs:
                run.font.size = Pt(text_size)
                run.font.color.rgb = textColor
                run.font.bold = bold
                run.font.italic = italic 
        tc = celda._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), bg_color_hex)
        tcPr.append(shd)

    for i, row_data in enumerate(conteo_rubros, start=1):
        row_cells = table.add_row().cells
        modificar_celda(celda=row_cells[0], texto=row_data["rubro__tipo"], text_size= 8, bg_color_hex=coloresT1[i % 2])
        modificar_celda(celda=row_cells[1], texto=row_data["total"], text_size= 8, bg_color_hex=coloresT1[i % 2])

    row_cells = table.add_row().cells
    modificar_celda(celda=row_cells[0], texto="Total", bg_color_hex="B38E5D", text_color_hex="FFFFFF", bold=True)
    modificar_celda(celda=row_cells[1], texto=sum(item['total'] for item in conteo_rubros), bg_color_hex="285C4D", text_color_hex="FFFFFF", bold=True)

    table1 = document.tables[1]
    for i, row_data in enumerate(datosA, start=1):
        row_cells = table1.add_row().cells
        modificar_celda(celda=row_cells[0], texto=row_data["rubro__tipo"], bg_color_hex=coloresT2[i % 2] )
        modificar_celda(celda=row_cells[1], texto=row_data["claveAcuerdo"], bg_color_hex=coloresT2[i % 2] )
        modificar_celda(celda=row_cells[2], texto=row_data["accionRTemporal__antecedente"], aliniar=2, bg_color_hex=coloresT2[i % 2] )
        modificar_celda(celda=row_cells[3], texto=row_data["accionRTemporal__descripcion"], aliniar=2, bg_color_hex=coloresT2[i % 2] )
        modificar_celda(celda=row_cells[4], texto=row_data["fecha_termino"], bg_color_hex=coloresT2[i % 2] )
        modificar_celda(celda=row_cells[5], texto=row_data["area__accionA2Temporal__area2__nickname"], bg_color_hex=coloresT2[i % 2] )

    final_stream = io.BytesIO()
    document.save(final_stream)
    final_stream.seek(0)
    
    return final_stream, context['oficina']

def generar_word(request):
    from datetime import datetime
    final_stream, oficina = _construir_documento_word()
    
    response = FileResponse(
        final_stream, 
        as_attachment=True, 
        filename=f"Instrucciones_{oficina}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )
    return response

def exportar_tabla():
    """
    Migra los datos de RegistroTemporal y AccionesTemporal a las tablas reales 
    (Registro y Acciones) y luego limpia las temporales.
    """
    from django.db import transaction

    with transaction.atomic():
        # 1. Migrar Registros
        registros_temp = RegistroTemporal.objects.all()
        mapeo_registros = {} # { id_temporal: objeto_real }

        for reg_temp in registros_temp:
            # Crear el registro real
            reg_real = Registro.objects.create(
                claveAcuerdo=reg_temp.claveAcuerdo,
                fecha_inicio=reg_temp.fecha_inicio,
                fecha_termino=reg_temp.fecha_termino,
                estado=reg_temp.estado,
                fecha_finalizacion=reg_temp.fecha_finalizacion,
                porcentaje_avance=reg_temp.porcentaje_avance,
                periodo=reg_temp.periodo
            )
            
            # Copiar ManyToMany de Rubros y Áreas
            reg_real.rubro.set(reg_temp.rubro.all())
            reg_real.area.set(reg_temp.area.all())
            
            # Guardar en el mapeo para las acciones
            mapeo_registros[reg_temp.idRegistro] = reg_real

        # 2. Migrar Acciones
        acciones_temp = AccionesTemporal.objects.all()
        for acc_temp in acciones_temp:
            # Crear la acción real
            acc_real = Acciones.objects.create(
                antecedente=acc_temp.antecedente,
                descripcion=acc_temp.descripcion
            )
            
            # Copiar ManyToMany de Áreas (area2)
            acc_real.area2.set(acc_temp.area2.all())
            
            # Copiar ManyToMany de Registros usando el mapeo
            registros_reales_vincular = []
            for r_temp in acc_temp.idRegistro.all():
                if r_temp.idRegistro in mapeo_registros:
                    # Usamos el objeto real creado anteriormente
                    registros_reales_vincular.append(mapeo_registros[r_temp.idRegistro])
            
            acc_real.idRegistro.set(registros_reales_vincular)

        # 3. Limpieza de tablas temporales
        # Primero acciones para evitar conflictos si hubiera FKs restrictivas (aunque aquí son M2M)
        AccionesTemporal.objects.all().delete()
        RegistroTemporal.objects.all().delete()

    return len(mapeo_registros)

@action(detail=False, methods=["get", "post"])
def guardar_exportar(request):
    from django.core.files.base import ContentFile
    from datetime import datetime
    
    try:
        # 1. Generar el documento Word
        final_stream, oficina = _construir_documento_word()
        nombre_archivo = f"Instrucciones_{oficina}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # 2. Guardar el archivo en el modelo ReporteGenerado
        reporte = ReporteGenerado()
        
        # Calcular clave a partir del primer RegistroTemporal
        reg_primer = RegistroTemporal.objects.first()
        if reg_primer:
            clave_acuerdo = reg_primer.claveAcuerdo
            # Formato: Primer digito + "00" + resto
            # Ejemplo: 001/CDMX/02/2026 -> 000/CDMX/02/2026
            if len(clave_acuerdo) >= 3:
                reporte.clave = clave_acuerdo[0] + "00" + clave_acuerdo[3:]
            else:
                reporte.clave = clave_acuerdo
        
        # Asignar el último periodo disponible
        ultimo_periodo = Periodo.objects.order_by('-id').first()
        if ultimo_periodo:
            reporte.periodo = ultimo_periodo

        reporte.archivo.save(nombre_archivo, ContentFile(final_stream.read()))
        
        # 3. Migrar los datos de las tablas temporales a las reales y limpiar
        num_migrados = exportar_tabla()
        
        messages.success(request, f"Archivo guardado y {num_migrados} registros migrados con éxito. Tablas temporales limpias.")
        return redirect('/supervision/')
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        messages.error(request, f"Ocurrió un error: {str(e)}")
        return redirect('/supervision/')